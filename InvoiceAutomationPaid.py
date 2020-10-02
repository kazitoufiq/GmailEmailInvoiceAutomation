
from docx import Document
from docx.shared import Inches
from datetime import datetime
from datetime import timedelta
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

document = Document('Invoice.docx')

InputRentalStart = '2020-10-04'




tables = document.tables


RentalStartYear =datetime.strptime(InputRentalStart, '%Y-%m-%d').strftime("%Y")
RentalStartMonth =datetime.strptime(InputRentalStart, '%Y-%m-%d').strftime("%m")
RentalStartDay =datetime.strptime(InputRentalStart, '%Y-%m-%d').strftime("%d")

tables[0].cell(0,0).text = 'Invoice No_ '+ RentalStartYear + RentalStartMonth + RentalStartDay

PaidInvoiceNo =  tables[0].cell(0, 0).text




tables[0].cell(0,1).text = "Issued: " +  datetime.today().strftime('%a, %d-%b-%Y')

tables[0].cell(0,1).text



DueBy ="Due by: " + (datetime.today() + timedelta(days=2)).strftime('%a, %d-%b-%Y')
    
tables[0].cell(0,2).text = DueBy 

DueBy



tables[1].cell(1,0).text



tables[3].cell(0,2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
tables[3].cell(0,2).text = '$290.00'
 



tables[3].cell(0,4).text ='$0.00'




StartDate=datetime.strptime(InputRentalStart, '%Y-%m-%d').strftime('%A, %d-%b-%Y')



EndDate=(datetime.strptime(InputRentalStart, '%Y-%m-%d') + 
         timedelta(days=6)).strftime('%A, %d-%b-%Y')



StartDate




EndDate



RentalPeriod = 'Van (Reg: 1OC9KC) Rental \nPeriod: '+ StartDate + ' to ' + EndDate +' (7 Days)'



RentalPeriod


tables[2].cell(0,1).text = RentalPeriod


run = tables[3].cell(0,7).add_paragraph().add_run("Payment Received. Thank you for your business!")
font = run.font
font.color.rgb = RGBColor(0x42, 0x24, 0xE9)




output_paid_invoice = PaidInvoiceNo + '_paid.docx'

document.save(output_paid_invoice)



from docx2pdf import convert

paid_invoice_pdf_file = PaidInvoiceNo + '_paid.pdf'

convert(output_paid_invoice,  paid_invoice_pdf_file)


